"""
Document Parser Utility

Extracts text content from various document formats (.txt, .pdf, .docx).
"""

import os
import io
from typing import Optional
import fitz  # PyMuPDF
from docx import Document
import PyPDF2
from app.core.logger import logger


class DocumentParser:
    """Parse text from various document formats"""
    
    @staticmethod
    def parse_txt(file_content: bytes) -> str:
        """Parse text from .txt file"""
        try:
            # Try UTF-8 first, fallback to latin-1
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1', errors='ignore')
        except Exception as e:
            logger.error(f"Error parsing TXT file: {e}")
            raise ValueError("Failed to parse TXT file")
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """Parse text from PDF file using PyMuPDF (primary) and PyPDF2 (fallback)"""
        try:
            # Try PyMuPDF first (better text extraction)
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                text_parts = []
                
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text = page.get_text()
                    text_parts.append(text)
                
                doc.close()
                
                full_text = "\n\n".join(text_parts).strip()
                
                if full_text:
                    return full_text
                else:
                    logger.warning("PyMuPDF extracted no text, trying PyPDF2")
            
            except Exception as fitz_error:
                logger.warning(f"PyMuPDF failed: {fitz_error}, trying PyPDF2")
            
            # Fallback to PyPDF2
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                text_parts.append(text)
            
            full_text = "\n\n".join(text_parts).strip()
            
            if not full_text:
                raise ValueError("No text could be extracted from PDF")
            
            return full_text
        
        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            raise ValueError("Failed to parse PDF file")
    
    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """Parse text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts).strip()
            
            if not full_text:
                raise ValueError("No text could be extracted from DOCX")
            
            return full_text
        
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise ValueError("Failed to parse DOCX file")
    
    @staticmethod
    def parse_file(file_content: bytes, filename: str) -> str:
        """
        Parse text from file based on extension.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename with extension
        
        Returns:
            Extracted text content
        
        Raises:
            ValueError: If file format is not supported or parsing fails
        """
        try:
            # Get file extension
            _, ext = os.path.splitext(filename.lower())
            
            if ext == '.txt':
                return DocumentParser.parse_txt(file_content)
            elif ext == '.pdf':
                return DocumentParser.parse_pdf(file_content)
            elif ext == '.docx':
                return DocumentParser.parse_docx(file_content)
            else:
                raise ValueError(f"Unsupported file format: {ext}. Supported formats: .txt, .pdf, .docx")
        
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error parsing file {filename}: {e}")
            raise ValueError(f"Failed to parse file: {str(e)}")


def parse_document(file_content: bytes, filename: str) -> str:
    """
    Helper function to parse document.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename with extension
    
    Returns:
        Extracted text content
    """
    return DocumentParser.parse_file(file_content, filename)
