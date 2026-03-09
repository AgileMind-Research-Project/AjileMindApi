"""
Report Export Utility

Export reports to PDF and DOCX formats with template styling.
"""

from typing import Dict, Any, Optional
import io
import base64
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image, Frame, PageTemplate
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.logger import logger
import json


class ReportExporter:
    """Export reports to various formats"""
    
    @staticmethod
    def _parse_color(color_str: str) -> tuple:
        """Parse hex color to RGB tuple"""
        try:
            color_str = color_str.lstrip('#')
            return tuple(int(color_str[i:i+2], 16) / 255.0 for i in (0, 2, 4))
        except:
            return (0, 0, 0)  # Default to black
    
    @staticmethod
    def _parse_color_docx(color_str: str) -> RGBColor:
        """Parse hex color to RGBColor for docx"""
        try:
            color_str = color_str.lstrip('#')
            r, g, b = tuple(int(color_str[i:i+2], 16) for i in (0, 2, 4))
            return RGBColor(r, g, b)
        except:
            return RGBColor(0, 0, 0)  # Default to black
    
    @staticmethod
    def _decode_base64_image(base64_string: str) -> Optional[io.BytesIO]:
        """
        Decode base64 image string to BytesIO.
        
        Args:
            base64_string: Base64 encoded image (may include data URI prefix)
        
        Returns:
            BytesIO object containing image data, or None if invalid
        """
        try:
            # Remove data URI prefix if present (e.g., "data:image/png;base64,")
            if ',' in base64_string:
                base64_string = base64_string.split(',', 1)[1]
            
            image_data = base64.b64decode(base64_string)
            return io.BytesIO(image_data)
        except Exception as e:
            logger.error(f"Error decoding base64 image: {e}")
            return None
    
    @staticmethod
    def export_to_pdf(
        report_data: Dict[str, Any],
        report_type: str,
        template_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export report to PDF format.
        
        Args:
            report_data: Report content dictionary
            report_type: Type of report (daily_standup, sprint_meeting, retrospective)
            template_config: Optional template configuration with header, footer, styles
        
        Returns:
            PDF file as bytes
        """
        try:
            buffer = io.BytesIO()
            
            # Prepare header/footer image data for page template
            header_img_data = None
            footer_img_data = None
            
            if report_data.get('header_image'):
                header_img_stream = ReportExporter._decode_base64_image(report_data['header_image'])
                if header_img_stream:
                    header_img_data = header_img_stream.getvalue()
            
            if report_data.get('footer_image'):
                footer_img_stream = ReportExporter._decode_base64_image(report_data['footer_image'])
                if footer_img_stream:
                    footer_img_data = footer_img_stream.getvalue()
            
            # Calculate margins based on header/footer images
            top_margin = 1.25*inch if header_img_data else 0.75*inch
            bottom_margin = 1.25*inch if footer_img_data else 0.75*inch
            
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=letter, 
                topMargin=top_margin, 
                bottomMargin=bottom_margin,
                leftMargin=1*inch,
                rightMargin=1*inch
            )
            
            # Create a function to draw header/footer on each page
            def add_header_footer(canvas_obj, doc_obj):
                canvas_obj.saveState()
                page_width, page_height = letter
                
                # Draw header image on every page (full width)
                if header_img_data:
                    try:
                        from reportlab.lib.utils import ImageReader
                        header_stream = io.BytesIO(header_img_data)
                        img_reader = ImageReader(header_stream)
                        # Position at top of page, full width edge to edge
                        img_width = page_width  # Full page width
                        img_height = 1.0*inch  # 1 inch height
                        x_pos = 0
                        y_pos = page_height - img_height
                        canvas_obj.drawImage(img_reader, x_pos, y_pos, width=img_width, height=img_height, preserveAspectRatio=False, mask='auto')
                    except Exception as e:
                        logger.error(f"Error drawing header image: {e}")
                
                # Draw footer image on every page (full width)
                if footer_img_data:
                    try:
                        from reportlab.lib.utils import ImageReader
                        footer_stream = io.BytesIO(footer_img_data)
                        img_reader = ImageReader(footer_stream)
                        # Position at bottom of page, full width edge to edge
                        img_width = page_width  # Full page width
                        img_height = 1.0*inch  # 1 inch height
                        x_pos = 0
                        y_pos = 0
                        canvas_obj.drawImage(img_reader, x_pos, y_pos, width=img_width, height=img_height, preserveAspectRatio=False, mask='auto')
                    except Exception as e:
                        logger.error(f"Error drawing footer image: {e}")
                
                canvas_obj.restoreState()
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Apply custom styles if provided
            if template_config and template_config.get('styles'):
                custom_styles = template_config['styles']
                
                # Update heading style
                if 'heading_color' in custom_styles:
                    color = ReportExporter._parse_color(custom_styles['heading_color'])
                    styles['Heading1'].textColor = colors.Color(*color)
                
                if 'font_size' in custom_styles:
                    styles['Normal'].fontSize = int(custom_styles['font_size'])
            
            story = []
            
            # Add header text if provided (from template_config)
            if template_config and template_config.get('header_content'):
                header = template_config['header_content']
                if header.get('text'):
                    alignment = TA_CENTER if header.get('alignment') == 'center' else TA_LEFT
                    header_style = ParagraphStyle(
                        'CustomHeader',
                        parent=styles['Heading1'],
                        alignment=alignment,
                        fontSize=int(header.get('font_size', 14)),
                        textColor=colors.Color(*ReportExporter._parse_color(header.get('color', '#000000')))
                    )
                    story.append(Paragraph(header['text'], header_style))
                    story.append(Spacer(1, 0.3*inch))
            
            # Add report title
            title = f"{report_type.replace('_', ' ').title()} Report"
            story.append(Paragraph(title, styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
            
            # Add report content based on type
            # Template-based reports store _sections_order; use generic renderer for them
            if '_sections_order' in report_data:
                ReportExporter._add_template_content(story, report_data, styles)
            elif report_type == 'daily_standup':
                ReportExporter._add_daily_standup_content(story, report_data, styles)
            elif report_type == 'sprint_meeting':
                ReportExporter._add_sprint_meeting_content(story, report_data, styles)
            elif report_type == 'retrospective':
                ReportExporter._add_retrospective_content(story, report_data, styles)
            elif report_type == 'brainstorming':
                ReportExporter._add_brainstorming_content(story, report_data, styles)
            
            # Add footer text if provided (from template_config)
            if template_config and template_config.get('footer_content'):
                footer = template_config['footer_content']
                if footer.get('text'):
                    story.append(Spacer(1, 0.5*inch))
                    alignment = TA_CENTER if footer.get('alignment') == 'center' else TA_LEFT
                    footer_style = ParagraphStyle(
                        'CustomFooter',
                        parent=styles['Normal'],
                        alignment=alignment,
                        fontSize=int(footer.get('font_size', 10)),
                        textColor=colors.Color(*ReportExporter._parse_color(footer.get('color', '#666666')))
                    )
                    story.append(Paragraph(footer['text'], footer_style))
            
            # Build PDF with header/footer callback
            doc.build(story, onFirstPage=add_header_footer, onLaterPages=add_header_footer)
            
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
        
        except Exception as e:
            logger.error(f"Error exporting report to PDF: {e}")
            raise
    
    @staticmethod
    def _add_template_content(story, data, styles):
        """Generic PDF renderer for template-based reports using _sections_order metadata"""
        sections_order = data.get('_sections_order', [])
        sorted_sections = sorted(sections_order, key=lambda s: s.get('order', 0))

        for section in sorted_sections:
            key = section.get('key', '')
            title = section.get('title', key.replace('_', ' ').title())
            sec_type = section.get('type', 'paragraph')
            value = data.get(key)

            if value is None:
                continue

            story.append(Paragraph(f"{title}:", styles['Heading2']))
            story.append(Spacer(1, 0.1 * inch))

            if sec_type in ('bullet_list', 'numbered_list'):
                items = value if isinstance(value, list) else [value]
                for item in items:
                    story.append(Paragraph(f"• {item}", styles['Normal']))

            elif sec_type == 'table' and isinstance(value, list) and value:
                # Build table from array of dicts or strings
                if isinstance(value[0], dict):
                    columns = list(value[0].keys())
                    table_data = [[col.replace('_', ' ').title() for col in columns]]
                    for row in value:
                        table_data.append([str(row.get(col, '')) for col in columns])
                else:
                    table_data = [['#', 'Item']]
                    for i, item in enumerate(value, 1):
                        table_data.append([str(i), str(item)])

                col_count = len(table_data[0])
                available_width = 6.5 * inch
                col_width = available_width / col_count
                table = Table(table_data, colWidths=[col_width] * col_count)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)

            else:
                text = str(value) if not isinstance(value, list) else ', '.join(str(v) for v in value)
                story.append(Paragraph(text, styles['Normal']))

            story.append(Spacer(1, 0.2 * inch))

    @staticmethod
    def _add_daily_standup_content(story, data, styles):
        """Add daily standup content to PDF – developer-centric format with tables"""
        team_updates = data.get('team_updates', [])

        if team_updates:
            # Create table for team updates
            table_data = [['Developer', 'Yesterday', 'Today', 'Blockers']]

            for dev in team_updates:
                if not isinstance(dev, dict):
                    continue
                name = dev.get('name', 'Unknown')
                role = dev.get('role', '')
                developer_name = f"{name}" + (f"\n({role})" if role else "")

                yesterday = dev.get('yesterday_tasks', [])
                yesterday_text = '\n'.join([f"• {task}" for task in yesterday]) if yesterday else 'No tasks'

                today = dev.get('today_tasks', [])
                today_text = '\n'.join([f"• {task}" for task in today]) if today else 'No tasks'

                blockers = dev.get('blockers', [])
                blockers_text = '\n'.join([f"• {b}" for b in blockers]) if blockers else 'No blockers'

                table_data.append([developer_name, yesterday_text, today_text, blockers_text])

            if len(table_data) > 1:  # Only add table if there are data rows
                table = Table(table_data, colWidths=[2*inch, 2.5*inch, 2.5*inch, 2.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.3*inch))
        else:
            # Legacy flat format fallback
            story.append(Paragraph("Yesterday's Work:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            for item in (data.get('yesterday_work') or []):
                text = item if isinstance(item, str) else str(item)
                story.append(Paragraph(f"• {text}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

            story.append(Paragraph("Today's Plan:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            for item in (data.get('today_plan') or []):
                text = item if isinstance(item, str) else str(item)
                story.append(Paragraph(f"• {text}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))

            story.append(Paragraph("Blockers & Issues:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            blockers = data.get('blockers') or []
            if blockers:
                for item in blockers:
                    text = item if isinstance(item, str) else str(item)
                    story.append(Paragraph(f"• {text}", styles['Normal']))
            else:
                story.append(Paragraph("No blockers reported.", styles['Normal']))

        # Blockers summary table
        blockers_summary = data.get('blockers_summary', [])
        if blockers_summary:
            story.append(Spacer(1, 0.15*inch))
            story.append(Paragraph("Blockers Summary:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))

            table_data = [['Title', 'Description', 'Reported By', 'Impact']]
            for bs in blockers_summary:
                if isinstance(bs, dict):
                    title = bs.get('title', 'Untitled')
                    desc = bs.get('description', '')
                    reported = ', '.join(bs.get('reported_by', []))
                    impact = bs.get('impact', '')
                    table_data.append([title, desc, reported, impact])

            if len(table_data) > 1:  # Only add table if there are data rows
                table = Table(table_data, colWidths=[2*inch, 2.5*inch, 1.5*inch, 2*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ]))
                story.append(table)
    
    @staticmethod
    def _add_sprint_meeting_content(story, data, styles):
        """Add sprint meeting content to PDF"""
        # Sprint goals
        story.append(Paragraph("Sprint Goals:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        for item in (data.get('sprint_goals') or []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Progress summary
        story.append(Paragraph("Progress Summary:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        # Handle string or list for progress_summary just in case
        prog = data.get('progress_summary') or []
        if isinstance(prog, str):
            prog = [prog]
        for item in prog:
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Issues & risks
        story.append(Paragraph("Issues & Risks:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        for item in (data.get('issues_risks') or []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Action items as table
        story.append(Paragraph("Action Items:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        action_items = data.get('action_items') or []
        if action_items:
            table_data = [['Task', 'Assignee', 'Due Date']]
            for item in action_items:
                task = item.get('task') or item.get('action') or 'Unknown Task'
                assignee = item.get('assignee', 'Unassigned')
                due_date = item.get('due_date', 'No deadline')
                table_data.append([task, assignee, due_date])

            table = Table(table_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No action items.", styles['Normal']))
    
    @staticmethod
    def _add_retrospective_content(story, data, styles):
        """Add retrospective content to PDF"""
        # What went well
        story.append(Paragraph("What Went Well:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        for item in (data.get('what_went_well') or []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # What didn't go well
        story.append(Paragraph("What Didn't Go Well:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        for item in (data.get('what_didnt_go_well') or []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Improvements
        story.append(Paragraph("Improvements:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        for item in (data.get('improvements') or []):
            story.append(Paragraph(f"• {item}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Action points as table
        story.append(Paragraph("Action Points:", styles['Heading2']))
        story.append(Spacer(1, 0.1*inch))
        action_points = data.get('action_points') or []
        if action_points:
            table_data = [['Action', 'Assignee']]
            for item in action_points:
                if isinstance(item, dict):
                    action = item.get('task') or item.get('action') or 'Unknown'
                    assignee = item.get('assignee', 'Unassigned')
                    table_data.append([action, assignee])
                else:
                    table_data.append([str(item), 'Unassigned'])

            table = Table(table_data, colWidths=[4*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No action points.", styles['Normal']))
    
    @staticmethod
    def export_to_docx(
        report_data: Dict[str, Any],
        report_type: str,
        template_config: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Export report to DOCX format.
        
        Args:
            report_data: Report content dictionary
            report_type: Type of report
            template_config: Optional template configuration
        
        Returns:
            DOCX file as bytes
        """
        try:
            doc = Document()
            
            # Get the default section
            section = doc.sections[0]
            
            # Set section margins for header/footer to allow full width images
            from docx.shared import Cm, Emu
            section.header_distance = Cm(0)
            section.footer_distance = Cm(0)
            
            # Get page dimensions
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            
            # Add header image to document header (appears on every page, full width)
            if report_data.get('header_image'):
                header_img_stream = ReportExporter._decode_base64_image(report_data['header_image'])
                if header_img_stream:
                    try:
                        doc_header = section.header
                        doc_header.is_linked_to_previous = False
                        # Clear existing paragraphs
                        for para in doc_header.paragraphs:
                            para.clear()
                        header_para = doc_header.paragraphs[0] if doc_header.paragraphs else doc_header.add_paragraph()
                        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Remove paragraph spacing and set negative indents to extend full width
                        header_para.paragraph_format.space_before = Pt(0)
                        header_para.paragraph_format.space_after = Pt(0)
                        header_para.paragraph_format.left_indent = -left_margin
                        header_para.paragraph_format.right_indent = -right_margin
                        run = header_para.add_run()
                        # Use full page width with compact height
                        run.add_picture(header_img_stream, width=page_width, height=Inches(1))
                    except Exception as e:
                        logger.error(f"Error adding header image to DOCX: {e}")
            
            # Add header text if provided (from template_config) - to document body, not header
            if template_config and template_config.get('header_content'):
                header = template_config['header_content']
                if header.get('text'):
                    p = doc.add_paragraph(header['text'])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if header.get('alignment') == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                    p.runs[0].font.size = Pt(int(header.get('font_size', 14)))
                    p.runs[0].font.color.rgb = ReportExporter._parse_color_docx(header.get('color', '#000000'))
                    doc.add_paragraph()
            
            # Add title
            title = doc.add_heading(f"{report_type.replace('_', ' ').title()} Report", 0)
            
            # Add report content based on type
            # Template-based reports store _sections_order; use generic renderer for them
            if '_sections_order' in report_data:
                ReportExporter._add_template_content_docx(doc, report_data)
            elif report_type == 'daily_standup':
                ReportExporter._add_daily_standup_content_docx(doc, report_data)
            elif report_type == 'sprint_meeting':
                ReportExporter._add_sprint_meeting_content_docx(doc, report_data)
            elif report_type == 'retrospective':
                ReportExporter._add_retrospective_content_docx(doc, report_data)
            elif report_type == 'brainstorming':
                ReportExporter._add_brainstorming_content_docx(doc, report_data)
            
            # Add footer text if provided (from template_config) - to document body
            if template_config and template_config.get('footer_content'):
                footer = template_config['footer_content']
                if footer.get('text'):
                    doc.add_paragraph()
                    p = doc.add_paragraph(footer['text'])
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if footer.get('alignment') == 'center' else WD_ALIGN_PARAGRAPH.LEFT
                    p.runs[0].font.size = Pt(int(footer.get('font_size', 10)))
                    p.runs[0].font.color.rgb = ReportExporter._parse_color_docx(footer.get('color', '#666666'))
            
            # Add footer image to document footer (appears on every page, full width)
            if report_data.get('footer_image'):
                footer_img_stream = ReportExporter._decode_base64_image(report_data['footer_image'])
                if footer_img_stream:
                    try:
                        doc_footer = section.footer
                        doc_footer.is_linked_to_previous = False
                        # Clear existing paragraphs
                        for para in doc_footer.paragraphs:
                            para.clear()
                        footer_para = doc_footer.paragraphs[0] if doc_footer.paragraphs else doc_footer.add_paragraph()
                        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Remove paragraph spacing and set negative indents to extend full width
                        footer_para.paragraph_format.space_before = Pt(0)
                        footer_para.paragraph_format.space_after = Pt(0)
                        footer_para.paragraph_format.left_indent = -left_margin
                        footer_para.paragraph_format.right_indent = -right_margin
                        run = footer_para.add_run()
                        # Use full page width with compact height
                        run.add_picture(footer_img_stream, width=page_width, height=Inches(1))
                    except Exception as e:
                        logger.error(f"Error adding footer image to DOCX: {e}")
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            docx_bytes = buffer.getvalue()
            buffer.close()
            
            return docx_bytes
        
        except Exception as e:
            logger.error(f"Error exporting report to DOCX: {e}")
            raise
    
    @staticmethod
    def _add_template_content_docx(doc, data):
        """Generic DOCX renderer for template-based reports using _sections_order metadata"""
        sections_order = data.get('_sections_order', [])
        sorted_sections = sorted(sections_order, key=lambda s: s.get('order', 0))

        for section in sorted_sections:
            key = section.get('key', '')
            title = section.get('title', key.replace('_', ' ').title())
            sec_type = section.get('type', 'paragraph')
            value = data.get(key)

            if value is None:
                continue

            doc.add_heading(f"{title}:", level=2)

            if sec_type in ('bullet_list', 'numbered_list'):
                items = value if isinstance(value, list) else [value]
                for item in items:
                    doc.add_paragraph(str(item), style='List Bullet')

            elif sec_type == 'table' and isinstance(value, list) and value:
                # Build table from array of dicts or strings
                if isinstance(value[0], dict):
                    columns = list(value[0].keys())
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(columns):
                        hdr_cells[i].text = col.replace('_', ' ').title()
                        for para in hdr_cells[i].paragraphs:
                            for run in para.runs:
                                run.bold = True
                    for row in value:
                        row_cells = table.add_row().cells
                        for i, col in enumerate(columns):
                            row_cells[i].text = str(row.get(col, ''))
                else:
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '#'
                    hdr_cells[1].text = 'Item'
                    for para in hdr_cells[0].paragraphs:
                        for run in para.runs:
                            run.bold = True
                    for para in hdr_cells[1].paragraphs:
                        for run in para.runs:
                            run.bold = True
                    for i, item in enumerate(value, 1):
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(i)
                        row_cells[1].text = str(item)
                doc.add_paragraph()

            else:
                text = str(value) if not isinstance(value, list) else ', '.join(str(v) for v in value)
                doc.add_paragraph(text)

    @staticmethod
    def _add_daily_standup_content_docx(doc, data):
        """Add daily standup content to DOCX – developer-centric format with tables"""
        team_updates = data.get('team_updates', [])

        if team_updates:
            # Create table for team updates
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Developer'
            header_cells[1].text = 'Yesterday'
            header_cells[2].text = 'Today'
            header_cells[3].text = 'Blockers'

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for dev in team_updates:
                if not isinstance(dev, dict):
                    continue

                row_cells = table.add_row().cells
                name = dev.get('name', 'Unknown')
                role = dev.get('role', '')
                developer_name = f"{name}" + (f"\n({role})" if role else "")
                row_cells[0].text = developer_name

                yesterday = dev.get('yesterday_tasks', [])
                yesterday_text = '\n'.join([f"• {task}" for task in yesterday]) if yesterday else 'No tasks'
                row_cells[1].text = yesterday_text

                today = dev.get('today_tasks', [])
                today_text = '\n'.join([f"• {task}" for task in today]) if today else 'No tasks'
                row_cells[2].text = today_text

                blockers = dev.get('blockers', [])
                blockers_text = '\n'.join([f"• {b}" for b in blockers]) if blockers else 'No blockers'
                row_cells[3].text = blockers_text

            doc.add_paragraph()  # Add spacing after table
        else:
            # Legacy fallback
            doc.add_heading("Yesterday's Work:", level=2)
            for item in (data.get('yesterday_work') or []):
                text = item if isinstance(item, str) else str(item)
                doc.add_paragraph(text, style='List Bullet')

            doc.add_heading("Today's Plan:", level=2)
            for item in (data.get('today_plan') or []):
                text = item if isinstance(item, str) else str(item)
                doc.add_paragraph(text, style='List Bullet')

            doc.add_heading("Blockers & Issues:", level=2)
            blockers = data.get('blockers') or []
            if blockers:
                for item in blockers:
                    text = item if isinstance(item, str) else str(item)
                    doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph("No blockers reported.")

        # Blockers summary table
        blockers_summary = data.get('blockers_summary', [])
        if blockers_summary:
            doc.add_heading("Blockers Summary:", level=2)

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Title'
            header_cells[1].text = 'Description'
            header_cells[2].text = 'Reported By'
            header_cells[3].text = 'Impact'

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for bs in blockers_summary:
                if isinstance(bs, dict):
                    row_cells = table.add_row().cells
                    row_cells[0].text = bs.get('title', 'Untitled')
                    row_cells[1].text = bs.get('description', '')
                    row_cells[2].text = ', '.join(bs.get('reported_by', []))
                    row_cells[3].text = bs.get('impact', '')
    
    @staticmethod
    def _add_sprint_meeting_content_docx(doc, data):
        """Add sprint meeting content to DOCX"""
        doc.add_heading("Sprint Goals:", level=2)
        for item in (data.get('sprint_goals') or []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Progress Summary:", level=2)
        # Handle potential string vs list
        prog = data.get('progress_summary') or []
        if isinstance(prog, str):
            prog = [prog]
        for item in prog:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Issues & Risks:", level=2)
        for item in (data.get('issues_risks') or []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Action Items:", level=2)
        action_items = data.get('action_items') or []
        if action_items:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Task'
            header_cells[1].text = 'Assignee'
            header_cells[2].text = 'Due Date'

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for item in action_items:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('task') or item.get('action') or 'Unknown Task'
                row_cells[1].text = item.get('assignee', 'Unassigned')
                row_cells[2].text = item.get('due_date', 'No deadline')

            doc.add_paragraph()  # Add spacing after table
        else:
            doc.add_paragraph("No action items.")
    
    @staticmethod
    def _add_retrospective_content_docx(doc, data):
        """Add retrospective content to DOCX"""
        doc.add_heading("What Went Well:", level=2)
        for item in (data.get('what_went_well') or []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("What Didn't Go Well:", level=2)
        for item in (data.get('what_didnt_go_well') or []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Improvements:", level=2)
        for item in (data.get('improvements') or []):
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading("Action Points:", level=2)
        action_points = data.get('action_points') or []
        if action_points:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'

            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Action'
            header_cells[1].text = 'Assignee'

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Add data rows
            for item in action_points:
                row_cells = table.add_row().cells
                if isinstance(item, dict):
                    row_cells[0].text = item.get('task') or item.get('action') or 'Unknown'
                    row_cells[1].text = item.get('assignee', 'Unassigned')
                else:
                    row_cells[0].text = str(item)
                    row_cells[1].text = 'Unassigned'

            doc.add_paragraph()  # Add spacing after table
        else:
            doc.add_paragraph("No action points.")

    @staticmethod
    def _add_brainstorming_content(story, data, styles):
        """Add brainstorming content to PDF"""
        # Meeting Topic and Objective
        if data.get('meeting_topic'):
            story.append(Paragraph(f"Topic: {data['meeting_topic']}", styles['Heading2']))
        if data.get('meeting_objective'):
            story.append(Paragraph(f"Objective: {data['meeting_objective']}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Summary
        if data.get('summary'):
            story.append(Paragraph("Summary:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(data['summary'], styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Participants
        if data.get('participants'):
            story.append(Paragraph("Participants:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(", ".join(data['participants']), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Top Ideas
        if data.get('top_ideas'):
            story.append(Paragraph("Top Ideas:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            for item in data['top_ideas']:
                story.append(Paragraph(f"★ {item}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # All Ideas Generated as table
        if data.get('ideas_generated'):
            story.append(Paragraph("Ideas Generated:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            ideas = data['ideas_generated']
            if ideas:
                table_data = [['Idea', 'Proposed By', 'Category']]
                for item in ideas:
                    if isinstance(item, dict):
                        idea = item.get('idea', '')
                        proposed_by = item.get('proposed_by', '')
                        category = item.get('category', '')
                        table_data.append([idea, proposed_by, category])
                    else:
                        table_data.append([str(item), '', ''])

                table = Table(table_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)
            story.append(Spacer(1, 0.2*inch))
        
        # Key Themes
        if data.get('key_themes'):
            story.append(Paragraph("Key Themes:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            for item in data['key_themes']:
                story.append(Paragraph(f"• {item}", styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        # Decisions Made as table
        if data.get('decisions_made'):
            story.append(Paragraph("Decisions Made:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            decisions = data['decisions_made']
            if decisions:
                table_data = [['Decision', 'Details']]
                for item in decisions:
                    if isinstance(item, dict):
                        decision = item.get('decision', '')
                        details = item.get('details', '')
                        table_data.append([decision, details])
                    else:
                        table_data.append([str(item), ''])

                table = Table(table_data, colWidths=[3*inch, 3*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)
            story.append(Spacer(1, 0.2*inch))
        
        # Next Steps as table
        if data.get('next_steps'):
            story.append(Paragraph("Next Steps:", styles['Heading2']))
            story.append(Spacer(1, 0.1*inch))
            steps = data['next_steps']
            if steps:
                table_data = [['Task', 'Assignee', 'Due Date', 'Priority']]
                for item in steps:
                    if isinstance(item, dict):
                        task = item.get('task', '')
                        assignee = item.get('assignee', '')
                        due_date = item.get('due_date', '')
                        priority = item.get('priority', '')
                        table_data.append([task, assignee, due_date, priority])
                    else:
                        table_data.append([str(item), '', '', ''])

                table = Table(table_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch, 1*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 10),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)

    @staticmethod
    def _add_brainstorming_content_docx(doc, data):
        """Add brainstorming content to DOCX"""
        # Meeting Topic and Objective
        if data.get('meeting_topic'):
            doc.add_heading(f"Topic: {data['meeting_topic']}", level=2)
        if data.get('meeting_objective'):
            doc.add_paragraph(f"Objective: {data['meeting_objective']}")
        
        # Summary
        if data.get('summary'):
            doc.add_heading("Summary:", level=2)
            doc.add_paragraph(data['summary'])
        
        # Participants
        if data.get('participants'):
            doc.add_heading("Participants:", level=2)
            doc.add_paragraph(", ".join(data['participants']))
        
        # Top Ideas
        if data.get('top_ideas'):
            doc.add_heading("Top Ideas:", level=2)
            for item in data['top_ideas']:
                doc.add_paragraph(f"★ {item}", style='List Bullet')
        
        # All Ideas Generated as table
        if data.get('ideas_generated'):
            doc.add_heading("Ideas Generated:", level=2)
            ideas = data['ideas_generated']
            if ideas:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Add header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Idea'
                hdr_cells[1].text = 'Proposed By'
                hdr_cells[2].text = 'Category'
                
                # Make header bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for item in ideas:
                    row_cells = table.add_row().cells
                    if isinstance(item, dict):
                        row_cells[0].text = item.get('idea', '')
                        row_cells[1].text = item.get('proposed_by', '')
                        row_cells[2].text = item.get('category', '')
                    else:
                        row_cells[0].text = str(item)
                        row_cells[1].text = ''
                        row_cells[2].text = ''
        
        # Key Themes
        if data.get('key_themes'):
            doc.add_heading("Key Themes:", level=2)
            for item in data['key_themes']:
                doc.add_paragraph(item, style='List Bullet')
        
        # Decisions Made as table
        if data.get('decisions_made'):
            doc.add_heading("Decisions Made:", level=2)
            decisions = data['decisions_made']
            if decisions:
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Decision'
                hdr_cells[1].text = 'Details'
                
                # Make header bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for item in decisions:
                    row_cells = table.add_row().cells
                    if isinstance(item, dict):
                        row_cells[0].text = item.get('decision', '')
                        row_cells[1].text = item.get('details', '')
                    else:
                        row_cells[0].text = str(item)
                        row_cells[1].text = ''
        
        # Next Steps as table
        if data.get('next_steps'):
            doc.add_heading("Next Steps:", level=2)
            steps = data['next_steps']
            if steps:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Add header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Task'
                hdr_cells[1].text = 'Assignee'
                hdr_cells[2].text = 'Due Date'
                hdr_cells[3].text = 'Priority'
                
                # Make header bold
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for item in steps:
                    row_cells = table.add_row().cells
                    if isinstance(item, dict):
                        row_cells[0].text = item.get('task', '')
                        row_cells[1].text = item.get('assignee', '')
                        row_cells[2].text = item.get('due_date', '')
                        row_cells[3].text = item.get('priority', '')
                    else:
                        row_cells[0].text = str(item)
                        row_cells[1].text = ''
                        row_cells[2].text = ''
                        row_cells[3].text = ''


def export_report(
    report_data: Dict[str, Any],
    report_type: str,
    export_format: str,
    template_config: Optional[Dict[str, Any]] = None
) -> bytes:
    """
    Helper function to export report.
    
    Args:
        report_data: Report content dictionary
        report_type: Type of report
        export_format: 'pdf' or 'docx'
        template_config: Optional template configuration
    
    Returns:
        Exported file as bytes
    """
    if export_format.lower() == 'pdf':
        return ReportExporter.export_to_pdf(report_data, report_type, template_config)
    elif export_format.lower() == 'docx':
        return ReportExporter.export_to_docx(report_data, report_type, template_config)
    else:
        raise ValueError(f"Unsupported export format: {export_format}")
