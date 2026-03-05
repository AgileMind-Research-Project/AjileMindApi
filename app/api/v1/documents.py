"""
Document Routes

API endpoints for document management and RAG-based chatbot.
Provides endpoints for fetching documents, dates, and generating chatbot responses.
"""

from fastapi import APIRouter, Depends, HTTPException, Query, File, UploadFile, Form
from fastapi.responses import JSONResponse
from datetime import date
from typing import List, Optional

from app.schemas.document import (
    DocumentCreate, DocumentResponse, DocumentListResponse,
    DocumentContentResponse, DocumentDateResponse,
    ChatQueryRequest, ChatQueryResponse, DocumentSearchRequest,
    DocumentSearchResponse, MultiDocChatRequest, MultiDocChatResponse
)
from app.services.document_service import document_service
from app.services.rag_service import get_rag_service  # Lazily create RAG service on first use
from app.core.logger import logger
from app.utils.jwt import get_current_user_from_token
from app.db.database import get_db, Database
from fastapi import status

# Create router
router = APIRouter(
    tags=["documents"],
    responses={404: {"description": "Not found"}},
)

# ==================== Document Management Endpoints ====================

@router.get(
    "/",
    response_model=List[DocumentResponse],
    summary="Get all documents",
    description="Fetch all documents for the current user's tenant, optionally filtered by date"
)
async def get_all_documents(
    limit: int = Query(1000, ge=1, le=10000),
    date: Optional[str] = Query(None, description="Filter by date (YYYY-MM-DD format)"),
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Get all documents
    
    - **limit**: Maximum number of results (default: 1000)
    - **date**: Optional date filter (YYYY-MM-DD format)
    
    Returns list of all documents for the user's tenant
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        if date:
            # Filter by specific date
            documents = await document_service.get_documents_by_date_str(schema, date, limit)
            logger.info(f"Fetched {len(documents)} documents for date {date} for user {current_user.get('id')}")
        else:
            documents = await document_service.get_all_documents(schema, limit)
            logger.info(f"Fetched {len(documents)} documents for user {current_user.get('id')}")
        
        return documents
        
        return documents
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error fetching documents"
        )


@router.post(
    "/",
    response_model=DocumentResponse,
    status_code=201,
    summary="Create a new document",
    description="Create a new document for use in RAG chatbot"
)
async def create_document(
    doc_data: DocumentCreate,
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Create a new document
    
    - **doc_title**: Document title (required)
    - **doc_content**: Full document content (required)
    - **uploaded_date**: Date of upload (required)
    - **category**: Optional category for classification
    """
    try:
        # Get schema from current user's tenant
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        document = await document_service.create_document(doc_data, schema)
        
        logger.info(f"Document created by user {current_user.get('id')}: {document.id}")
        
        return document
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error creating document: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error creating document"
        )


@router.get(
    "/dates",
    response_model=List[DocumentDateResponse],
    summary="Get unique document dates",
    description="Fetch all unique uploaded dates with document counts for date picker"
)
async def get_unique_dates(
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Get all unique document dates
    
    Returns list of dates when documents were uploaded with counts
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        dates = await document_service.get_unique_dates(schema)
        
        logger.info(f"Fetched unique document dates for user {current_user.get('id')}")
        
        return dates
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching unique dates: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error fetching unique dates"
        )


@router.get(
    "/by-date/{uploaded_date}",
    response_model=List[DocumentListResponse],
    summary="Get documents by date",
    description="Fetch document list for a specific upload date"
)
async def get_documents_by_date(
    uploaded_date: date,
    limit: int = Query(100, ge=1, le=1000),
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Get documents for a specific date
    
    - **uploaded_date**: Date to filter documents (YYYY-MM-DD format)
    - **limit**: Maximum number of results (default: 100)
    
    Returns list of documents uploaded on the specified date
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        documents = await document_service.get_documents_by_date(
            uploaded_date,
            schema,
            limit
        )
        
        logger.info(
            f"Fetched {len(documents)} documents for date {uploaded_date} "
            f"for user {current_user.get('id')}"
        )
        
        return documents
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching documents by date: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error fetching documents"
        )


@router.get(
    "/{document_id}",
    response_model=DocumentResponse,
    summary="Get document by ID",
    description="Retrieve complete document with all details"
)
async def get_document(
    document_id: int,
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Get a specific document by ID
    
    Returns complete document including content
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        document = await document_service.get_document_by_id(document_id, schema)
        
        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Document not found: {document_id}"
            )
        
        return document
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching document: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error fetching document"
        )


@router.get(
    "/{document_id}/content",
    response_model=DocumentContentResponse,
    summary="Get document content",
    description="Retrieve document content for RAG processing"
)
async def get_document_content(
    document_id: int,
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Get document content
    
    Returns document with full content for RAG chatbot
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        document = await document_service.get_document_content(document_id, schema)
        
        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Document not found: {document_id}"
            )
        
        return document
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching document content: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error fetching document content"
        )


# ==================== RAG Chatbot Endpoints ====================

@router.post(
    "/chat",
    response_model=ChatQueryResponse,
    summary="Chat with document using RAG",
    description="Send a query and get response based on selected document content. If document_id is not provided and search_all is True, searches all documents."
)
async def chat_with_document(
    query: ChatQueryRequest,
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Query the chatbot for a specific document or search all documents
    
    - **document_id**: ID of the document to use as context (optional if search_all is True)
    - **query**: The user's question
    - **search_all**: If True, search across all documents
    
    Returns chatbot response based on document content
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        # If search_all is True or document_id is not provided, search all documents
        if query.search_all or query.document_id is None:
            # Use multi-document search, optionally filtered by date
            documents = await document_service.get_all_documents_content(
                schema, 
                limit=100, 
                filter_date=query.filter_date
            )
            
            if not documents:
                from datetime import datetime as dt
                return ChatQueryResponse(
                    document_id=0,
                    document_title="No Documents",
                    user_query=query.query,
                    chatbot_response="No documents available to search. Please upload documents first.",
                    timestamp=dt.utcnow(),
                    source_document=None
                )
            
            # Generate response using RAG with multiple documents
            rag = get_rag_service()
            multi_response = await rag.generate_response_from_multiple_documents(documents, query.query)
            
            logger.info(
                f"Multi-doc chatbot response generated, searched {multi_response.searched_documents} documents "
                f"for user {current_user.get('id')}"
            )
            
            # Convert to ChatQueryResponse
            return ChatQueryResponse(
                document_id=multi_response.document_id,
                document_title=multi_response.document_title,
                user_query=multi_response.user_query,
                chatbot_response=multi_response.chatbot_response,
                timestamp=multi_response.timestamp,
                source_document=f"Found in: {multi_response.document_title} (searched {multi_response.searched_documents} documents)"
            )
        
        # Single document search
        document = await document_service.get_document_content(query.document_id, schema)
        
        if not document:
            raise HTTPException(
                status_code=404,
                detail=f"Document not found: {query.document_id}"
            )
        
        # Validate document has content
        if not document.doc_content or len(document.doc_content.strip()) < 10:
            logger.warning(f"Document {query.document_id} has no content or too short")
            # Return a helpful response instead of error
            from datetime import datetime as dt
            return ChatQueryResponse(
                document_id=document.id,
                document_title=document.doc_title or "Unknown",
                user_query=query.query,
                chatbot_response=f"The document '{document.doc_title}' appears to be empty or has insufficient content for analysis. Please upload a document with text content.",
                timestamp=dt.utcnow()
            )
        
        logger.debug(f"Processing chat for document {query.document_id}, content length: {len(document.doc_content)}")
        
        # Generate response using RAG (initialize service lazily)
        rag = get_rag_service()
        response = await rag.generate_response(document, query.query)
        
        logger.info(
            f"Chatbot response generated for document {query.document_id} "
            f"by user {current_user.get('id')}"
        )
        
        return response
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error generating chatbot response: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error generating chatbot response"
        )


@router.post(
    "/chat/all",
    response_model=MultiDocChatResponse,
    summary="Chat across all documents using RAG",
    description="Search across all documents to find the most relevant one and answer the query"
)
async def chat_with_all_documents(
    query: MultiDocChatRequest,
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Query the chatbot across all documents
    
    - **query**: The user's question
    - **uploaded_date**: Optional date filter (None = search all dates)
    
    Returns chatbot response from the most relevant document
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        # Fetch all documents (optionally filtered by date)
        if query.uploaded_date:
            # Get documents for specific date but with content
            documents = await document_service.get_documents_by_date_with_content(
                query.uploaded_date, schema, limit=100
            )
        else:
            # Get all documents
            documents = await document_service.get_all_documents_content(schema, limit=100)
        
        if not documents:
            from datetime import datetime
            return MultiDocChatResponse(
                document_id=0,
                document_title="No Documents",
                user_query=query.query,
                chatbot_response="No documents available to search. Please upload documents first.",
                timestamp=datetime.utcnow(),
                relevance_score=0.0,
                searched_documents=0
            )
        
        # Generate response using RAG with multiple documents
        rag = get_rag_service()
        response = await rag.generate_response_from_multiple_documents(documents, query.query)
        
        logger.info(
            f"Multi-doc chatbot response generated, searched {response.searched_documents} documents "
            f"for user {current_user.get('id')}"
        )
        
        return response
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error generating multi-doc chatbot response: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error generating chatbot response"
        )


# ==================== Search Endpoints ====================

@router.post(
    "/search",
    response_model=DocumentSearchResponse,
    summary="Search documents",
    description="Search documents by title, content, date, or category"
)
async def search_documents(
    search_req: DocumentSearchRequest,
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Search documents
    
    - **query**: Search keywords
    - **uploaded_date**: Optional date filter
    - **category**: Optional category filter
    - **limit**: Maximum results (1-100, default 10)
    
    Returns matching documents
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        results = await document_service.search_documents(
            search_req.query,
            schema,
            search_req.uploaded_date,
            search_req.category,
            search_req.limit
        )
        
        logger.info(
            f"Document search executed: {search_req.query} "
            f"(found {len(results)} results)"
        )
        
        return DocumentSearchResponse(
            total=len(results),
            results=results
        )
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error searching documents: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error searching documents"
        )


@router.get(
    "/category/{category}",
    response_model=List[DocumentListResponse],
    summary="Get documents by category",
    description="Fetch all documents in a specific category"
)
async def get_documents_by_category(
    category: str,
    limit: int = Query(100, ge=1, le=1000),
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Get documents by category
    
    - **category**: Category name
    - **limit**: Maximum number of results
    
    Returns all documents in the specified category
    """
    try:
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        documents = await document_service.get_documents_by_category(
            category,
            schema,
            limit
        )
        
        logger.info(
            f"Fetched {len(documents)} documents in category '{category}' "
            f"for user {current_user.get('id')}"
        )
        
        return documents
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error fetching documents by category: {e}")
        raise HTTPException(
            status_code=500,
            detail="Error fetching documents"
        )


# ==================== Health Check ====================

@router.post(
    "/upload",
    response_model=DocumentResponse,
    status_code=201,
    summary="Upload a document file",
    description="Upload a PDF, TXT, or DOCX file for RAG chatbot"
)
async def upload_document(
    file: UploadFile = File(...),
    doc_title: str = Form(...),
    category: str = Form(default="general"),
    current_user: dict = Depends(get_current_user_from_token)
):
    """
    Upload a document file
    
    - **file**: Document file (PDF, TXT, or DOCX)
    - **doc_title**: Document title (required)
    - **category**: Optional category for classification
    """
    try:
        # Get schema from current user's tenant
        schema = current_user.get("schema") or current_user.get("tenant_schema")
        
        if not schema:
            raise HTTPException(
                status_code=400,
                detail="User schema not found"
            )
        
        # Validate file type
        valid_types = ['application/pdf', 'text/plain', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        if file.content_type not in valid_types:
            raise HTTPException(
                status_code=400,
                detail="Invalid file type. Please upload PDF, TXT, or DOCX files only."
            )
        
        # Read file content
        content = await file.read()
        doc_content = ""
        
        # Extract text based on file type
        if file.content_type == 'application/pdf':
            # Extract text from PDF
            doc_content = await extract_text_from_pdf(content)
        elif file.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Extract text from DOCX
            doc_content = await extract_text_from_docx(content)
        else:
            # Plain text file
            doc_content = content.decode('utf-8', errors='ignore')
        
        # Validate extracted content
        if not doc_content or len(doc_content.strip()) < 10:
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from the uploaded file. Please ensure the file contains readable text."
            )
        
        logger.info(f"Extracted {len(doc_content)} characters from uploaded file: {file.filename}")
        
        # Create document using the file upload data
        doc_data = DocumentCreate(
            doc_title=doc_title,
            doc_content=doc_content,
            uploaded_date=date.today(),
            category=category
        )
        
        document = await document_service.create_document(doc_data, schema)
        
        logger.info(f"Document uploaded by user {current_user.get('id')}: {document.id}")
        
        return document
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error uploading document: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Error uploading document: {str(e)}"
        )


async def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF file bytes using PyMuPDF (fitz)"""
    logger.info(f"Attempting to extract text from PDF, content size: {len(content)} bytes")
    
    # Try PyMuPDF first
    try:
        import fitz  # PyMuPDF
        logger.info("Using PyMuPDF (fitz) for PDF extraction")
        
        # Open PDF from bytes
        doc = fitz.open(stream=content, filetype="pdf")
        num_pages = len(doc)
        logger.info(f"PDF opened successfully, {num_pages} pages found")
        
        text_parts = []
        for page_num in range(num_pages):
            page = doc[page_num]
            text = page.get_text()
            logger.debug(f"Page {page_num + 1}: extracted {len(text)} chars")
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        full_text = "\n\n".join(text_parts)
        logger.info(f"PyMuPDF: Extracted {len(full_text)} characters from {num_pages} pages")
        
        if full_text.strip():
            return full_text
        else:
            logger.warning("PyMuPDF extracted empty text, trying PyPDF2...")
            
    except ImportError as e:
        logger.warning(f"PyMuPDF not available: {e}")
    except Exception as e:
        logger.error(f"PyMuPDF extraction error: {e}", exc_info=True)
    
    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        import io
        logger.info("Using PyPDF2 for PDF extraction")
        
        reader = PdfReader(io.BytesIO(content))
        num_pages = len(reader.pages)
        logger.info(f"PDF opened with PyPDF2, {num_pages} pages found")
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            logger.debug(f"Page {i + 1}: extracted {len(text) if text else 0} chars")
            if text and text.strip():
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"PyPDF2: Extracted {len(full_text)} characters from {num_pages} pages")
        return full_text
        
    except ImportError as e:
        logger.error(f"PyPDF2 not available: {e}")
        return ""
    except Exception as e:
        logger.error(f"PyPDF2 extraction error: {e}", exc_info=True)
        return ""


async def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX file bytes"""
    try:
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except ImportError:
        logger.error("python-docx not installed. Cannot extract DOCX text.")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


@router.delete(
    "/{document_id}",
    summary="Delete document",
    description="Delete a report by ID"
)
async def delete_document(
    document_id: int,
    current_user: dict = Depends(get_current_user_from_token),
    db: Database = Depends(get_db)
):
    """Delete a report"""
    try:
        # Get tenant schema from user
        tenant_schema = current_user.get("tenant_schema") or current_user.get("schema", "gmail")
        
        # Check if report exists
        query = f"""
            SELECT id FROM {tenant_schema}.reports 
            WHERE id = %s
        """
        result = await db.execute_query(query, (document_id,), fetch_one=True)
        
        if not result:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Report with ID {document_id} not found"
            )
        
        # Soft delete by changing status to draft
        delete_query = f"""
            UPDATE {tenant_schema}.reports 
            SET status = 'draft'
            WHERE id = %s
        """
        await db.execute_query(delete_query, (document_id,), commit=True)
        
        logger.info(f"Report {document_id} deleted successfully")
        
        return {
            "message": "Report deleted successfully",
            "document_id": document_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting document: {str(e)}"
        )


@router.get(
    "/health",
    summary="Document service health check",
    description="Check if document service is running"
)
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "document-rag-chatbot",
        "timestamp": str(date.today())
    }
